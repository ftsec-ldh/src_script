import os,time
from info.api import get_company,aiqicha_get
from filter.socket_getIP import domain_to_ip
from docx import Document
from docx.shared import Inches,Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def cnvd_src_page(domain, leak_type, leak_url):

    description = {"信息泄露": "信息泄露可能会导致黑客进一步利用敏感信息盗取更关键的内容，甚至导致系统产生RCE漏洞",
                   "逻辑缺陷": "逻辑漏洞就是指攻击者利用业务的设计缺陷，获取敏感信息或破坏业务的完整性",
                   "SQL注入": "攻击者未经授权可以访问数据库中的数据",
                   "XSS": "反射xss通过引诱用户点击一个链接到目标网站的恶意链接来实施攻击,dom型则更具危害性，受害者正常进入页面则会遭受攻击",
                   "弱口令": "弱口令可能导致未经授权的用户获得对系统、应用程序或数据的访问权限。这可能会导致数据泄露、篡改、删除或者系统被滥用,如果用户在多个网站或服务中使用相同的弱密码，一旦其中一个网站的密码泄露，攻击者就能够使用该密码尝试登录其他网站，进而导致更大范围的信息泄露",
                   "命令执行": "远程代码执行简称RCE，是一类软件安全缺陷/漏洞。RCE 漏洞将允许恶意行为人通过 LAN、WAN 或 Internet 在远程计算机上执行自己选择的任何代码，属于更广泛的任意代码执行 (ACE) 漏洞类别。"

    }

    suggestions = {
        "逻辑缺陷": "修复不严谨的业务判断逻辑，对关键参数和用户COOKIE或鉴权认证进行校验",
        "SQL注入": "严格限制变量类型，比如整型变量就采用intval()函数过滤，数据库中的存储字段必须对应为int型",
        "XSS": "白名单过滤 根据白名单的标签和属性对数据进行过滤，以此来对可执行的脚本进行清除(如script标签，img标签的onerror属性等",
        "信息泄露": "设置访问该信息的权限配置，或只允许某一特定IP访问",
        "弱口令": "使用足够长度、包含大小写字母、数字和特殊字符的密码；还可以实施密码策略，要求用户设置符合安全标准的密码，如长度、复杂度要求，并对密码进行定期的过期和强制更改",
        "命令执行": "1.通用的修复方案，升级插件/框架/服务最新版\n2.如若必须使用危险函数，那么针对危险函数进行过滤\n3.在进入执行命令函数前进行严格的检测和过滤\n4.尽量不要使用命令执行函数，不能完全控制的危险函数最好不使用，如果非要用的话可以加验证防止被其他人利用"
    }


    if os.path.exists("cnvd_cookies.txt"):

        if leak_type == "弱口令":
            weak_account = input("请输入弱口令账号：")
            weak_password = input("请输入弱口令密码：")

        print("正在使用爱站查询域名信息，请耐心等待...")
        company_name = get_company(domain,1)

        print("正在使用爱企查查询企业信息，请耐心等待...")
        company_info = aiqicha_get(company_name, 1)

        ip = domain_to_ip(domain)
        province = company_info['province']
        city = company_info['city']

        if leak_type == "弱口令":
            doc = Document()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f'{company_name}网站存在{leak_type}\n\n弱口令账号：{weak_account}\n弱口令密码：{weak_password}')
            run.font.name = '等线'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
            doc.add_picture(f"aizhan_{company_name}.png", width=Inches(8.0))
            doc.add_picture(f"aiqicha_{company_name}.png", width=Inches(8.0))
            doc.save(f'{company_name}存在{leak_type}.docx')
        else:
            doc = Document()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f'{company_name}网站存在{leak_type}')
            run.font.name = '等线'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
            doc.add_picture(f"aizhan_{company_name}.png", width=Inches(8.0))
            doc.add_picture(f"aiqicha_{company_name}.png", width=Inches(8.0))
            doc.save(f'{company_name}存在{leak_type}.docx')

        from selenium import webdriver
        from selenium.webdriver.chrome.service import Service
        from selenium.webdriver.common.by import By
        from selenium.webdriver.support.wait import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.webdriver.common.action_chains import ActionChains
        from selenium.webdriver.support.ui import Select
        s = Service('drivers/win64/chromedriver.exe')
        options = webdriver.ChromeOptions()
        # 禁用GPU
        options.add_argument('--disable-gpu')
        options.add_argument('--no-sandbox')
        options.add_argument('window-size=1920,1080')

        # 去除顶端浏览器自动化测试提示
        options.add_experimental_option("excludeSwitches", ["enable-automation"])

        options.add_argument("--disable-blink-features=AutomationControlled")

        options.add_argument("--disable-blink-features")
        # 禁用自动化扩展程序，绕过限制
        options.add_experimental_option('useAutomationExtension', False)


        cnvd_driver = webdriver.Chrome(service=s,options=options)

        cnvd_driver.get("https://www.cnvd.org.cn/")
        WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH, "//a[@href='/user/login' and text()='登录']")))
        time.sleep(2)
        cnvd_driver.delete_all_cookies()

        with open("cnvd_cookies.txt","r+") as cookie_file:
            cookies = eval(cookie_file.read())
        for cookies in cookies:
            cnvd_driver.add_cookie(cookies)


        cnvd_driver.get("https://www.cnvd.org.cn/flaw/create")

        actions = ActionChains(cnvd_driver)
        try:
            element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH, "//input[@name='unitName']")))#涉事单位
        except Exception:
            print("cookie到期，请重启软件重新登录！")
            exit()
        actions.send_keys_to_element(element,company_name).perform()
        
        element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH, "//span[@id='select2-param_city-container']")))#所在省份-点击去掉记录
        actions.click(element).perform()

        element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH, "//input[@name='flowIP']")))#所属IP
        element.send_keys(ip)

        element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH, "//span[@id='select2-param_province-container']")))#所在省份
        actions.click(element).perform()

        element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH, f"//li[contains(text(), '{province}')]")))#省选项
        actions.click(element).perform()

        element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH, f"//span[@id='select2-param_city-container']")))#所在城市
        actions.click(element).perform()

        element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH, f"//li[contains(text(),'{city}')]")))#城市选项
        actions.click(element).perform()


        element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_all_elements_located((By.XPATH, f"//input[@placeholder='XXXX网站/XXXX系统/XXXX平台/XXXX']")))#漏洞网站公司名
        element[0].send_keys(f"{company_name}网站")

        s1 = Select(cnvd_driver.find_element(By.NAME, "titlel"))#漏洞类型
        s1.select_by_visible_text(leak_type)

        if leak_type == "弱口令":
            element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH,f"//input[@name='weakPasswordAccountNumber']")))#弱口令账号
            element.send_keys(weak_account)

            element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH, f"//input[@name='weakPasswordPwd']")))#弱口令密码
            element.send_keys(weak_password)


        element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH, f"//textarea[@placeholder='只能填写漏洞涉及的url，并且url仅允许以http或https开头，多个url请用英文逗号(,)隔开']")))#漏洞URL
        element.send_keys(leak_url)

        element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH, f"//textarea[@name='descriptionShow']")))#漏洞描述
        element.send_keys(f"{description[leak_type]}")

        element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH, f"//textarea[@name='tempWay']")))#临时解决方案
        element.send_keys(f"{suggestions[leak_type]}")

        element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH, f"//input[@id='flawAttFile']")))#文件上传
        element.send_keys(f"{os.getcwd()}\{company_name}存在{leak_type}.docx")


        #用户中心
        input()
    else:
        print("未检测到cookie")
        login()


def login():
    account = input("邮件：")
    password = input("密码：")
    from selenium import webdriver
    from selenium.webdriver.chrome.service import Service
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.wait import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    s = Service('drivers/win64/chromedriver.exe')
    options = webdriver.ChromeOptions()
    options.add_argument("--disable-blink-features=AutomationControlled")  # 绕过cnvd检测
    cnvd_driver = webdriver.Chrome(service=s, options=options)

    cnvd_driver.get("https://www.cnvd.org.cn/user/login")
    element = WebDriverWait(cnvd_driver, 10).until(EC.presence_of_element_located((By.XPATH, "//input[@id='email']")))
    element.send_keys(account)

    element = cnvd_driver.find_element(By.XPATH, "//input[@id='password']")
    element.send_keys(password)
    print("请手动输入验证码并登录")
    WebDriverWait(cnvd_driver, 600).until(EC.presence_of_element_located((By.XPATH, "//a[@id='rightMenuli4' and text()='我的任务']")))
    time.sleep(5)
    cookies = cnvd_driver.get_cookies()
    with open('cnvd_cookies.txt', 'w+') as output:
        output.write(str(cookies))
    print("保存cookie完毕，请重新执行本软件")
    exit()
